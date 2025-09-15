#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import io
import tempfile
import subprocess

import streamlit as st
from PyPDF2 import PdfReader, PdfWriter
from pdf2image import convert_from_path
from PIL import Image, ImageFilter, ImageOps, ImageEnhance
import pytesseract

# optional imports (graceful fallback)
try:
    from pdf2docx import Converter
    HAS_PDF2DOCX = True
except Exception:
    HAS_PDF2DOCX = False

from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ---------------- UI setup ----------------
st.set_page_config(page_title="PDF → Word (Kurdish-Optimized OCR)", page_icon="📄", layout="centered")
st.title("📄 PDF → Word Converter — Kurdish Optimized (Sorani/Kurmanji)")
st.caption("Convert any PDF (including scans) to editable Word. Best support for Sorani **ckb** and Kurmanji **kmr**.")

with st.expander("ℹ️ Tips for best Kurdish OCR"):
    st.markdown(
        "- Use Tesseract language codes: **ckb** (Sorani, Arabic script), **kmr** (Kurmanji, Latin script).  \n"
        "- For mixed content add **ara** and **eng**, e.g. `ckb+ara+eng`.  \n"
        "- If OCRmyPDF isn’t installed, the app falls back to a Kurdish-aware OCR pipeline."
    )

# ---------------- Sidebar ----------------
st.sidebar.header("OCR Settings")
lang_options = [
    "ckb", "kmr", "ara", "eng",
    "ckb+ara", "ckb+eng", "ckb+ara+eng",
    "kmr+eng", "kmr+ara+eng"
]
langs = st.sidebar.selectbox("Tesseract languages", lang_options, index=lang_options.index("ckb+ara+eng"))
psm = st.sidebar.selectbox(
    "PSM (Page Segmentation Mode)",
    ["3 - Fully automatic", "4 - Single column", "6 - Uniform block", "11 - Sparse text", "12 - Sparse + OSD"],
    index=2
)
oem = st.sidebar.selectbox("OEM (Engine Mode)", ["3 - Default", "1 - LSTM only", "0 - Legacy only"], index=1)
try_ocrmypdf = st.sidebar.checkbox("Use OCRmyPDF first (best layout if available)", value=True)
dpi = st.sidebar.slider("Image DPI (fallback OCR)", 200, 500, 350, 50)
contrast = st.sidebar.slider("Contrast (fallback)", 1.0, 2.5, 1.4, 0.1)
sharpness = st.sidebar.slider("Sharpness (fallback)", 1.0, 2.5, 1.2, 0.1)
max_pages = st.sidebar.number_input("Max pages to process (0 = all)", min_value=0, value=0, step=1)

uploaded = st.file_uploader("Upload a PDF", type=["pdf"])

# ---------------- Helpers ----------------
def which(cmd):
    from shutil import which as _which
    return _which(cmd)

def has_text_layer(pdf_path: str) -> bool:
    try:
        reader = PdfReader(pdf_path)
        for page in reader.pages:
            txt = page.extract_text() or ""
            if txt.strip():
                return True
        return False
    except Exception:
        return False

def limited_pages(pdf_path: str, limit: int) -> str:
    if not limit or limit <= 0:
        return pdf_path
    reader = PdfReader(pdf_path)
    writer = PdfWriter()
    n = min(len(reader.pages), limit)
    for i in range(n):
        writer.add_page(reader.pages[i])
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    with open(tmp.name, "wb") as f:
        writer.write(f)
    return tmp.name

def tesseract_config(psm_sel: str, oem_sel: str, langs: str) -> str:
    psm_val = int(psm_sel.split(" - ")[0])
    oem_val = int(oem_sel.split(" - ")[0])
    cfg = f"--oem {oem_val} --psm {psm_val} -c preserve_interword_spaces=1"
    # Minor tweak for Arabic-script layout
    if "ckb" in langs and "ara" in langs:
        cfg += " -c textord_old_xheight=1"
    return cfg

def preprocess(img: Image.Image, contrast=1.4, sharpness=1.2) -> Image.Image:
    g = ImageOps.grayscale(img)
    g = ImageEnhance.Contrast(g).enhance(contrast)
    g = ImageEnhance.Sharpness(g).enhance(sharpness)
    g = g.filter(ImageFilter.MedianFilter(size=3))
    g = ImageOps.autocontrast(g)
    return g

def add_paragraph(doc: Document, text: str, rtl: bool = False):
    p = doc.add_paragraph(text)
    if rtl:
        pPr = p._p.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), "1")
        pPr.append(bidi)

def parse_tsv(tsv_text: str):
    lines = (tsv_text or "").splitlines()
    if not lines:
        return {}
    header = lines[0].split('\t')
    idx = {k: i for i, k in enumerate(header)}
    out = {}
    for row in lines[1:]:
        cols = row.split('\t')
        if len(cols) <= max(idx.values()):
            continue
        try:
            text = cols[idx['text']].strip()
            conf = float(cols[idx['conf']])
            if not text or conf < 0:
                continue
            block = int(cols[idx.get('block_num', 0)])
            line = int(cols[idx.get('line_num', 0)])
            left = int(cols[idx.get('left', 0)])
            top = int(cols[idx.get('top', 0)])
            key = (block, line)
            out.setdefault(key, []).append((left, top, text, conf))
        except Exception:
            continue
    for k in out:
        out[k].sort(key=lambda t: t[0])
    return out

def fallback_ocr_to_docx(pdf_path: str, out_docx: str, langs: str, dpi=350, page_limit=0, psm_sel="6 - Uniform block", oem_sel="1 - LSTM only", contrast=1.4, sharpness=1.2):
    images = convert_from_path(pdf_path, dpi=dpi)
    doc = Document()
    # Font choice: Arabic-script friendly if Sorani
    style = doc.styles['Normal']
    style.font.name = 'Noto Naskh Arabic' if 'ckb' in langs else 'Calibri'
    style.font.size = Pt(12)

    count = len(images) if page_limit in (0, None) else min(len(images), page_limit)
    cfg = tesseract_config(psm_sel, oem_sel, langs)
    rtl = ('ckb' in langs)

    for i in range(count):
        im = preprocess(images[i], contrast=contrast, sharpness=sharpness)
        tsv = pytesseract.image_to_data(im, lang=langs, config=cfg, output_type=pytesseract.Output.STRING)
        blocks = parse_tsv(tsv)

        if i > 0:
            doc.add_page_break()

        if blocks:
            keys_sorted = sorted(blocks.keys(), key=lambda k: min([t[1] for t in blocks[k]]))
            for key in keys_sorted:
                words = [w[2] for w in blocks[key]]
                line_text = " ".join(words).strip()
                if line_text:
                    add_paragraph(doc, line_text, rtl=rtl)
        else:
            # pure text fallback
            txt = pytesseract.image_to_string(im, lang=langs, config=cfg).strip()
            add_paragraph(doc, txt, rtl=rtl)

    doc.save(out_docx)

def run_ocrmypdf(in_pdf: str, out_pdf: str, langs: str) -> tuple[int, str]:
    cmd = [
        "ocrmypdf",
        "--skip-text",
        "-l", langs,
        "--rotate-pages", "--deskew",
        "--clean-final", "--remove-background",
        "--optimize", "3",
        "--output-type", "pdf",
        in_pdf, out_pdf
    ]
    res = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True)
    return res.returncode, res.stdout

def pdf_to_docx_direct(pdf_path: str, docx_path: str):
    if not HAS_PDF2DOCX:
        raise RuntimeError("pdf2docx not installed; cannot run direct conversion.")
    cv = Converter(pdf_path)
    try:
        cv.convert(docx_path, start=0, end=None)
    finally:
        cv.close()

# ---------------- Main flow ----------------
if uploaded:
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_in:
        tmp_in.write(uploaded.getvalue())
        pdf_in = tmp_in.name

    # limit for speed if user set max_pages
    pdf_work = limited_pages(pdf_in, max_pages)

    st.write("### Conversion mode")
    text_layer = has_text_layer(pdf_work)
    if text_layer:
        if HAS_PDF2DOCX:
            st.info("🧾 Detected a text layer → using **direct PDF→DOCX** for best layout.")
        else:
            st.warning("🧾 Detected a text layer, but **pdf2docx** is not installed. Falling back to OCR path.")

    c1, c2 = st.columns(2)
    with c1:
        go = st.button("🚀 Convert to Word")
    with c2:
        stop = st.button("❌ Cancel")

    if go:
        with st.spinner("Processing..."):
            out_name = os.path.splitext(uploaded.name)[0] + "_converted.docx"
            out_buf = io.BytesIO()

            try:
                # Path A: direct conversion if we have text layer and pdf2docx
                if text_layer and HAS_PDF2DOCX:
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_docx:
                        pdf_to_docx_direct(pdf_work, tmp_docx.name)
                        with open(tmp_docx.name, "rb") as f:
                            out_buf.write(f.read())
                    note = "Direct PDF→DOCX (text layer preserved)."

                # Path B: OCR flow
                else:
                    used_ocrmypdf = False
                    if try_ocrmypdf and which("ocrmypdf"):
                        st.write("🔍 Running OCRmyPDF (best layout if installed)...")
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_ocr_pdf:
                            code, log = run_ocrmypdf(pdf_work, tmp_ocr_pdf.name, langs=langs)
                        if code == 0:
                            used_ocrmypdf = True
                            if HAS_PDF2DOCX:
                                st.success("✅ OCRmyPDF succeeded. Converting searchable PDF to DOCX…")
                                with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_docx:
                                    pdf_to_docx_direct(tmp_ocr_pdf.name, tmp_docx.name)
                                    with open(tmp_docx.name, "rb") as f:
                                        out_buf.write(f.read())
                                note = "OCRmyPDF → DOCX (layout-friendly)."
                            else:
                                # If pdf2docx missing, still do fallback OCR on the OCRed PDF's images
                                st.warning("pdf2docx not installed; using fallback OCR builder.")
                                with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_docx:
                                    fallback_ocr_to_docx(tmp_ocr_pdf.name, tmp_docx.name, langs, dpi, max_pages, psm, oem, contrast, sharpness)
                                    with open(tmp_docx.name, "rb") as f:
                                        out_buf.write(f.read())
                                note = "OCRmyPDF + Fallback OCR builder."

                        else:
                            st.warning("OCRmyPDF failed or not available. Falling back to Kurdish-optimized OCR.")
                            st.caption(f"Log:\n\n```\n{log}\n```")

                    if not used_ocrmypdf:
                        st.write("🧠 Kurdish-optimized OCR fallback (preprocess + line reconstruction + RTL for ckb)…")
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_docx:
                            fallback_ocr_to_docx(
                                pdf_work, tmp_docx.name,
                                langs=langs, dpi=dpi, page_limit=max_pages,
                                psm_sel=psm, oem_sel=oem, contrast=contrast, sharpness=sharpness
                            )
                            with open(tmp_docx.name, "rb") as f:
                                out_buf.write(f.read())
                        note = "Fallback OCR (Kurdish-optimized)."

                st.success(f"Done! {note}")
                st.download_button(
                    "⬇️ Download .docx",
                    data=out_buf.getvalue(),
                    file_name=out_name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception as e:
                st.error(f"Conversion failed: {e}")

    if stop:
        st.stop()

st.markdown("---")
st.caption("Install Tesseract language packs: `ckb`, `kmr`, plus `ara`, `eng` if needed. If `pdf2docx` or `ocrmypdf` are missing, the app still works via the fallback OCR builder.")
